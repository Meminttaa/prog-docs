from docx import Document


def replace_text_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)

    for key, value in replacements.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", value)

    if full_text != "".join(run.text for run in paragraph.runs):
        paragraph.clear()
        paragraph.add_run(full_text)


def replace_in_container(paragraphs, replacements):
    for paragraph in paragraphs:
        replace_text_in_paragraph(paragraph, replacements)


def replace_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_container(cell.paragraphs, replacements)


def create_docx_for_student(template_path, output_path, student):
    doc = Document(template_path)

    # основной текст
    replace_in_container(doc.paragraphs, student)
    replace_in_tables(doc.tables, student)

    # колонтитулы (ВОТ ЧЕГО НЕ ХВАТАЛО)
    for section in doc.sections:
        replace_in_container(section.header.paragraphs, student)
        replace_in_tables(section.header.tables, student)

        replace_in_container(section.footer.paragraphs, student)
        replace_in_tables(section.footer.tables, student)

    doc.save(output_path)

